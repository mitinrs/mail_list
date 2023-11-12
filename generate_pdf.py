import openpyxl
import re
from docx import Document
from comtypes.client import CreateObject
import os
import sys

def clean_filename(filename):
    return re.sub(r'[\\/*?:"<>|]', '_', filename)

def find_placeholders(doc):
    placeholder_pattern = r"\[(.+?)\]"
    placeholders = set()

    for paragraph in doc.paragraphs:
        placeholders.update(re.findall(placeholder_pattern, paragraph.text))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                placeholders.update(re.findall(placeholder_pattern, cell.text))

    return placeholders

def docx_to_pdf(docx_filename, pdf_filename):
    try:
        word = CreateObject('Word.Application')
        doc = word.Documents.Open(docx_filename)
        doc.SaveAs(pdf_filename, 17)  # Используем порядковый номер для формата файла
        doc.Close()
        word.Quit()
    except Exception as e:
        print(f"Failed to save '{pdf_filename}'. Error: {e}")

if __name__ == '__main__':

    # Забираем аргументы командной строки
    if len(sys.argv) > 2:
        list_path = sys.argv[1]  # Используем имя excel файла, переданное из .bat файла
        template_path = sys.argv[2]  # Используем имя excel файла, переданное из .bat файла
        working_directory = sys.argv[3] # Используем путь к рабочей папке, переданный из .bat файла
    else:
        print("Ошибка: Неверный вызов вункции")
        sys.exit(1)

    # Пути к файлам и папкам
    excel_filename = os.path.join(working_directory, list_path)
    template_filename = os.path.join(working_directory, template_path)
    doc_folder = os.path.join(working_directory, 'DOC')
    pdf_folder = os.path.join(working_directory, 'PDF')

    # Создание папок для DOC и PDF файлов, если они еще не созданы
    os.makedirs(doc_folder, exist_ok=True)
    os.makedirs(pdf_folder, exist_ok=True)

    # Загрузка файла Excel
    wb = openpyxl.load_workbook(excel_filename)
    sheet = wb.active

    # Загрузка шаблона Word
    template_document = Document(template_filename)

    # Найти все плейсхолдеры в шаблоне
    placeholders = find_placeholders(template_document)

    # Создание словаря для сопоставления заголовков столбцов и плейсхолдеров
    column_titles = {cell.value: idx for idx, cell in enumerate(sheet[1])}

    # Обработка каждой строки в Excel
    for row_index, row in enumerate(sheet.iter_rows(min_row=2, values_only=True), start=2):
        doc = Document(template_filename)

        # Замена плейсхолдеров в параграфах и таблицах
        for placeholder in placeholders:
            placeholder_tag = f'[{placeholder}]'
            column_index = column_titles.get(placeholder)
            if column_index is not None:
                replacement_text = str(row[column_index] or '')
                for paragraph in doc.paragraphs:
                    paragraph.text = paragraph.text.replace(placeholder_tag, replacement_text)
                for table in doc.tables:
                    for table_row in table.rows:
                        for cell in table_row.cells:
                            cell.text = cell.text.replace(placeholder_tag, replacement_text)

        # Генерация имен файлов и путей
        outgoing_number = str(row[column_titles['Исходящий номер']])
        addressee_name = str(row[column_titles['Имя адресата']])
        filename = f'{outgoing_number}-{addressee_name}'
        cleaned_filename = clean_filename(filename)
        new_docx_path = os.path.join(doc_folder, f'{cleaned_filename}.docx')
        new_pdf_path = os.path.join(pdf_folder, f'{cleaned_filename}.pdf')

        # Сохранение документа Word
        try:
            doc.save(new_docx_path)
        except Exception as e:
            print(f"Failed to save '{new_docx_path}'. Error: {e}")

        # Конвертация в PDF
        try:
            docx_to_pdf(new_docx_path, new_pdf_path)
        except Exception as e:
            print(f"Failed to save '{new_pdf_path}'. Error: {e}")

        print(f"Processed row {row_index}: {new_docx_path} and {new_pdf_path}")
