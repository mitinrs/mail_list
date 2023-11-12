import openpyxl
import re
from docx import Document
from comtypes.client import CreateObject
import os
import sys

def clean_filename(filename):
    return re.sub(r'[\\/*?:"<>|]', '_', filename)

def get_cell_value(cell):
    # Если ячейка содержит формулу, вернуть результат формулы
    if cell.data_type == 'f':
        return cell.value
    # В противном случае, вернуть само значение
    return cell.internal_value

def find_placeholders(doc):
    placeholder_pattern = r"\[(.+?)\]"
    placeholders = set()

    for paragraph in doc.paragraphs:
        placeholders.update(re.findall(placeholder_pattern, paragraph.text))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                placeholders.update(re.findall(placeholder_pattern, cell.text))

    return placeholders

def docx_to_pdf(docx_filename, pdf_filename):
    try:
        word = CreateObject('Word.Application')
        doc = word.Documents.Open(docx_filename)
        doc.SaveAs(pdf_filename, 17)  # Используем порядковый номер для формата файла
        doc.Close()
        word.Quit()
    except Exception as e:
        print(f"Failed to save '{pdf_filename}'. Error: {e}")

def replace_placeholder(paragraph, placeholder, replacement):
    for run in paragraph.runs:
        run.text = run.text.replace(placeholder, replacement)

def replace_placeholder_in_table(table, placeholder, replacement):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_placeholder(paragraph, placeholder, replacement)


if __name__ == '__main__':

    # Забираем аргументы командной строки
    if len(sys.argv) > 2:
        list_path = sys.argv[1]  # Используем имя excel файла, переданное из .bat файла
        template_path = sys.argv[2]  # Используем имя excel файла, переданное из .bat файла
        working_directory = sys.argv[3] # Используем путь к рабочей папке, переданный из .bat файла
    else:
        # Пути к файлам и папкам
        working_directory = 'D:\\GIT\\email list'
        list_path = os.path.join(working_directory, 'list.xlsx')
        template_path = os.path.join(working_directory, 'template.docx')       
        #print("Ошибка: Неверный вызов вункции")
        #sys.exit(1)

    # Пути к файлам и папкам
    excel_filename = os.path.join(working_directory, list_path)
    template_filename = os.path.join(working_directory, template_path)
    doc_folder = os.path.join(working_directory, 'DOC')
    pdf_folder = os.path.join(working_directory, 'PDF')

    # Создание папок для DOC и PDF файлов, если они еще не созданы
    os.makedirs(doc_folder, exist_ok=True)
    os.makedirs(pdf_folder, exist_ok=True)

    # Загрузка файла Excel
    wb = openpyxl.load_workbook(excel_filename, data_only=True)
    sheet = wb.active

    # Загрузка шаблона Word
    template_document = Document(template_filename)

    # Найти все плейсхолдеры в шаблоне
    placeholders = find_placeholders(template_document)

    # Создание словаря для сопоставления заголовков столбцов и плейсхолдеров
    column_titles = {cell.value: idx for idx, cell in enumerate(sheet[1])}
    
    # Создание словаря для сопоставления плейсхолдеров и заголовков столбцов Excel
    column_titles = {cell.value: idx for idx, cell in enumerate(sheet[1], start=1)}
    placeholder_to_column = {f'[{placeholder}]': sheet.cell(row=1, column=column_titles[placeholder]).value for placeholder in placeholders if placeholder in column_titles}


    # Обработка каждой строки в Excel
    for row_index, row in enumerate(sheet.iter_rows(min_row=2), start=2):
        doc = Document(template_filename)

        # Замена плейсхолдеров в параграфах и таблицах, сохраняя форматирование
        for placeholder_tag, column_title in placeholder_to_column.items():
            column_index = column_titles[column_title]
            cell = row[column_index - 1]
            replacement_text = str(get_cell_value(cell) or '')
            for paragraph in doc.paragraphs:
                replace_placeholder(paragraph, placeholder_tag, replacement_text)
            for table in doc.tables:
                replace_placeholder_in_table(table, placeholder_tag, replacement_text)

        # Генерация имен файлов и путей
        outgoing_number_cell = row[column_titles['Исходящий номер'] - 1]
        addressee_name_cell = row[column_titles['Имя адресата'] - 1]
        outgoing_number = str(get_cell_value(outgoing_number_cell) or '')
        addressee_name = str(get_cell_value(addressee_name_cell) or '')
        filename = f'{outgoing_number}-{addressee_name}'
        cleaned_filename = clean_filename(filename)
        new_docx_path = os.path.join(doc_folder, f'{cleaned_filename}.docx')
        new_pdf_path = os.path.join(pdf_folder, f'{cleaned_filename}.pdf')

        # Сохранение документа Word
        try:
            doc.save(new_docx_path)
        except Exception as e:
            print(f"Failed to save '{new_docx_path}'. Error: {e}")

        # Конвертация в PDF
        try:
            docx_to_pdf(new_docx_path, new_pdf_path)
        except Exception as e:
            print(f"Failed to save '{new_pdf_path}'. Error: {e}")

        print(f"Processed row {row_index}: {new_docx_path} and {new_pdf_path}")
